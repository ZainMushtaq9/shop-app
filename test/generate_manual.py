  import os
  import glob
  from docx import Document
  from docx.shared import Pt, RGBColor, Inches
  from docx.enum.text import WD_ALIGN_PARAGRAPH
  from reportlab.lib.pagesizes import A4
  from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
  from reportlab.lib.colors import HexColor
  from reportlab.lib.units import inch
  from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    Image, HRFlowable, PageBreak
  )
  from reportlab.lib.enums import TA_CENTER, TA_LEFT
  from datetime import datetime

  # Directories
  os.makedirs('test/manual', exist_ok=True)
  os.makedirs('test/reports', exist_ok=True)

  SCREENSHOTS = sorted(glob.glob('test/screenshots/*.png'))
  TIMESTAMP   = datetime.now().strftime('%Y-%m-%d %H:%M')
  TEAL        = HexColor('#006D77')

  SECTIONS = [
    {
      'title': 'Super Business Shop — Mکمل Guide',
      'subtitle': 'Pakistani Shopkeepers ke Liye',
    },
    {
      'heading': '1. Pehli Baar Setup',
      'english': '1. First Time Setup',
      'content': '''
  App kholen aur "Naya Account Banao" tap karein.

  Step 1: Dukaan ka naam, qism, shehar, phone bharein
  Step 2: "Subah kitna cash tha?" — opening cash bharein
  Step 3: Pehla product add karein (skip kar sakte hain)
  Step 4: Zubaan chunein — Urdu ya English
  Step 5: Gmail backup setup (baad mein bhi kar sakte hain)
  Step 6: "Dukaan Shuru Karo" — Dashboard khul jayega

  TIP: Urdu choose karein for best experience.
      ''',
      'screenshot': '02_after_login_dashboard'
    },
    {
      'heading': '2. Bikri Kaise Karein (POS)',
      'english': '2. How to Make a Sale',
      'content': '''
  Left sidebar ya bottom nav mein "Bikri Karo" tap karein.

  Step 1: Product search karein ya barcode scan karein
  Step 2: Product tap karein — cart mein chala jayega
  Step 3: Quantity + ya - se badlein
  Step 4: Discount lagaein (optional)
  Step 5: Customer select karein (ya Walk-in rehne dein)
  Step 6: Cash / Udhaar / Partial chunein
  Step 7: "Bill Banao" tap karein
  Step 8: Bill preview mein Print ya WhatsApp karein

  IMPORTANT: Is screen par koi ad nahi hai.
  Aap ka kaam bilkul free flow mein hoga.
      ''',
      'screenshot': '09_pos_screen'
    },
    {
      'heading': '3. Maal (Inventory) Manage Karein',
      'english': '3. Manage Your Stock',
      'content': '''
  "Maal" section mein jayein.

  Filters:
  - Tamam: Sare products
  - Kam Stock: Stock khatam hone wala
  - Thanda Maal: 30+ din se nahi bika
  - Kam Munafa: 10% se kam margin

  Product Add Karein:
  1. + button tap karein
  2. Naam (English + Urdu) bharein
  3. Khareed aur bikri qeemat bharein
  4. Margin khud calculate ho jaye ga
  5. Stock quantity bharein
  6. Save tap karein

  TIP: Stock Value tab mein dekhin
  aap ke poore maal ki qeemat kya hai.
      ''',
      'screenshot': '08_inventory_list'
    },
    {
      'heading': '4. Udhaar Track Karein',
      'english': '4. Track Customer Credit',
      'content': '''
  "Grahak" section mein jayein.

  - Lal rang: Grahak ne aap ko paisa dena hai
  - Hara rang: Aap ne grahak ko paisa dena hai
  - Grahak tap karein — poori history dekhin
  - "Payment Record Karo" tap karein
  - "WhatsApp Reminder" bhejein

  Customer Portal:
  Grahak ko apna hisaab mobile par dekhne ke liye
  "QR Code" ya "Link Bhejo" use karein.
  Grahak link khol ke apna poora hisaab dekh sakta hai.
  Grahak kuch edit nahi kar sakta — sirf dekh sakta hai.
      ''',
      'screenshot': '10_customers_list'
    },
    {
      'heading': '5. Munafa Nuqsan (Profit/Loss)',
      'english': '5. Profit and Loss Report',
      'content': '''
  "Riportein" section mein jayein.

  Date Filter:
  - Aaj: Aaj ka hisaab
  - Hafte: Is hafte ka
  - Mahine: Is mahine ka
  - Custom: Apni taareekh chunein

  Hisaab kuch iss tarah hota hai:
  Kul Bikri - Maal ki Lagat - Kharche = Asal Munafa

  Hara card = Munafa (Acha hai!)
  Lal card  = Nuqsan (Dhyan dein!)

  PDF Download: "Download Karo" button
  WhatsApp:     "WhatsApp" button — report share karein
      ''',
      'screenshot': '11_reports_screen'
    },
    {
      'heading': '6. Roz ka Hisaab (Daily Cash)',
      'english': '6. Daily Cash Management',
      'content': '''
  "Aaj ka Hisaab" section mein jayein.

  Subah:
  - "Subah kitna cash tha?" — number bharein
  - "Shuru Karo" tap karein

  Din mein:
  - Live balance dekhtay rahein
  - Sab kuch automatic update hota hai

  Raat ko:
  - "Hisaab Band Karo" tap karein
  - Cash gino aur enter karein
  - System bataye ga farq hai ya nahi

  Hisaab theek = Mubarak ho!
  Farq = Reason likhein
      ''',
      'screenshot': '14_daily_cash_screen'
    },
    {
      'heading': '7. Internet Nahi? Koi baat nahi!',
      'english': '7. Works Without Internet',
      'content': '''
  App bilkul internet ke bina kaam karta hai!

  Internet ke bina kar sakte hain:
  ✅ Bikri kar sakte hain
  ✅ Bills bana sakte hain
  ✅ Stock check kar sakte hain
  ✅ Reports dekh sakte hain
  ✅ Kharche add kar sakte hain

  Kya hoga internet aane par:
  - Sab data apne aap sync ho jaye ga
  - "Data sync ho raha hai" message ayega
  - Customer portal bhi update ho jaye ga

  Orange banner = Internet nahi (kaam jaari hai)
  Green banner  = Internet aa gaya (sync ho raha hai)
      ''',
      'screenshot': '15_offline_banner_appears'
    },
    {
      'heading': '8. Data Backup (Gmail Drive)',
      'english': '8. Backup Your Data',
      'content': '''
  Settings mein "Data Save" section mein jayein.

  Ek baar setup:
  1. "Gmail se Connect Karo" tap karein
  2. Apni shop wali Gmail se login karein
  3. Permission dein — Done!

  Automatic backup:
  - Har raat 11 baje backup hota hai
  - Google Drive mein save hota hai
  - 30 backups rakhta hai

  Manual backup:
  - "Abhi Backup Lo" tap karein
  - 4 stages mein hota hai
  - Success message ayega

  Agar phone kho jaye ya kharab ho:
  - Naye phone par app install karein
  - "Data Wapas Lao" tap karein
  - Poora data wapas aa jaye ga
      ''',
      'screenshot': '17_backup_screen'
    },
    {
      'heading': '9. Ads aur Premium',
      'english': '9. Ads and Premium',
      'content': '''
  Free plan mein chota banner ad aata hai:
  - Sirf dashboard aur reports par
  - Bilkul top par — chhota 320x50
  - Bada X button — click karein — band ho jaye ga
  - POS par koi ad NAHI hai
  - Kisi bhi form par koi ad NAHI hai

  Pro Plan (PKR 999/month):
  ✅ Koi ad nahi
  ✅ Customer portal
  ✅ Unlimited products
  ✅ 2 co-helpers
  ✅ All reports

  Business Plan (PKR 1,999/month):
  ✅ Sab kuch Pro mein
  ✅ 5 co-helpers
  ✅ 2 dukaanen
  ✅ Bulk WhatsApp
      ''',
      'screenshot': '07_ad_banner_visible'
    },
  ]

  # ── GENERATE DOCX ─────────────────────────────────────
  def make_docx():
    doc = Document()

    # Title
    t = doc.add_heading(
      'Super Business Shop — Mukammal Guide', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
      run.font.color.rgb = RGBColor(0, 109, 119)

    doc.add_paragraph(
      f'Generated: {TIMESTAMP}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    for sec in SECTIONS[1:]:
      h = doc.add_heading(sec['heading'], level=1)
      for run in h.runs:
        run.font.color.rgb = RGBColor(0, 109, 119)

      for line in sec['content'].strip().split('\\n'):
        p = doc.add_paragraph(line.strip())
        if line.strip().startswith(('✅','❌','Step','TIP','IMPORTANT')):
          for run in p.runs:
            run.bold = True

      # Add screenshot if exists
      ss = sec.get('screenshot', '')
      if ss:
        matches = [s for s in SCREENSHOTS if ss in s]
        if matches:
          try:
            doc.add_picture(matches[0], width=Inches(5))
          except Exception:
            pass
      doc.add_paragraph('')

    doc.save('test/manual/user_manual.docx')
    print('✅ DOCX: test/manual/user_manual.docx')

  # ── GENERATE PDF ──────────────────────────────────────
  def make_pdf():
    doc = SimpleDocTemplate(
      'test/manual/user_manual.pdf',
      pagesize=A4,
      rightMargin=inch,
      leftMargin=inch,
      topMargin=inch,
      bottomMargin=inch,
    )

    styles = getSampleStyleSheet()
    title_s = ParagraphStyle('T',
      fontSize=22, textColor=TEAL,
      alignment=TA_CENTER, spaceAfter=8,
      fontName='Helvetica-Bold')
    h_s = ParagraphStyle('H',
      fontSize=14, textColor=TEAL,
      spaceBefore=16, spaceAfter=6,
      fontName='Helvetica-Bold')
    body_s = ParagraphStyle('B',
      fontSize=10, spaceAfter=4,
      fontName='Helvetica', leading=14)
    sub_s = ParagraphStyle('S',
      fontSize=11, alignment=TA_CENTER,
      textColor=HexColor('#555555'), spaceAfter=4)

    story = []
    story.append(Paragraph(
      'Super Business Shop — Mukammal Guide', title_s))
    story.append(Paragraph(
      'Pakistani Shopkeepers ke Liye', sub_s))
    story.append(Paragraph(
      f'Generated: {TIMESTAMP}', sub_s))
    story.append(Spacer(1, 0.2*inch))
    story.append(HRFlowable(
      width='100%', color=TEAL, thickness=2))
    story.append(Spacer(1, 0.2*inch))

    for sec in SECTIONS[1:]:
      story.append(Paragraph(sec['heading'], h_s))
      story.append(HRFlowable(
        width='100%', color=HexColor('#CCCCCC'),
        thickness=0.5))
      story.append(Spacer(1, 0.1*inch))

      for line in sec['content'].strip().split('\\n'):
        line = line.strip()
        if not line:
          story.append(Spacer(1, 0.05*inch))
          continue
        clean = (line
          .replace('&', '&amp;')
          .replace('<', '&lt;')
          .replace('>', '&gt;'))
        if line.startswith(('✅','❌','Step','TIP','IMPORTANT')):
          story.append(Paragraph(
            f'<b>{clean}</b>', body_s))
        else:
          story.append(Paragraph(clean, body_s))

      ss = sec.get('screenshot', '')
      if ss:
        matches = [s for s in SCREENSHOTS if ss in s]
        if matches:
          try:
            story.append(Spacer(1, 0.1*inch))
            story.append(Image(matches[0],
              width=5*inch, height=3*inch))
          except Exception:
            pass
      story.append(Spacer(1, 0.2*inch))

    doc.build(story)
    print('✅ PDF: test/manual/user_manual.pdf')

  # ── GENERATE TXT ──────────────────────────────────────
  def make_txt():
    lines = [
      '=' * 60,
      'SUPER BUSINESS SHOP — MUKAMMAL GUIDE',
      f'Generated: {TIMESTAMP}',
      '=' * 60, '',
    ]
    for sec in SECTIONS[1:]:
      lines.append(sec['heading'])
      lines.append('-' * 40)
      lines.append(sec['content'].strip())
      lines.append('')

    with open('test/manual/user_manual.txt', 'w',
              encoding='utf-8') as f:
      f.write('\\n'.join(lines))
    print('✅ TXT: test/manual/user_manual.txt')

  # ── GENERATE HTML MANUAL ──────────────────────────────
  def make_html():
    ss_tags = ''
    for s in SCREENSHOTS[:30]:
      name = os.path.basename(s)
      ss_tags += f'''
      <div class="ss-card">
        <img src="../screenshots/{name}" loading="lazy"
             onclick="this.requestFullscreen()">
        <p>{name.replace(".png","").replace("_"," ")}</p>
      </div>'''

    video_tags = ''
    videos = glob.glob('test/videos/*.webm')
    for v in sorted(videos):
      name = os.path.basename(v).replace('.webm','')
      video_tags += f'''
      <div class="v-card">
        <h3>{name.replace("_"," ")}</h3>
        <video controls>
          <source src="../videos/{os.path.basename(v)}"
                  type="video/webm">
        </video>
      </div>'''

    html = f'''<!DOCTYPE html>
  <html lang="ur" dir="rtl">
  <head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width">
    <title>Super Business Shop — Guide</title>
    <style>
      * {{ box-sizing: border-box; margin: 0; padding: 0; }}
      body {{ font-family: sans-serif;
               background: #f5f5f5; }}
      .header {{ background: #006D77; color: white;
                  padding: 30px; text-align: center; }}
      .header h1 {{ font-size: 28px; }}
      .section {{ max-width: 1000px; margin: 20px auto;
                   background: white; border-radius: 12px;
                   padding: 24px;
                   box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
      h2 {{ color: #006D77; margin-bottom: 16px; }}
      pre {{ background: #f9f9f9; padding: 16px;
              border-radius: 8px; white-space: pre-wrap;
              font-family: sans-serif; line-height: 1.7; }}
      .ss-grid {{ display: grid;
                   grid-template-columns: repeat(3,1fr);
                   gap: 12px; margin-top: 16px; }}
      .ss-card img {{ width: 100%; border-radius: 8px;
                       cursor: pointer; border: 1px solid #eee; }}
      .ss-card p {{ font-size: 11px; color: #666;
                     margin-top: 4px; text-align: center; }}
      .v-grid {{ display: grid;
                  grid-template-columns: repeat(2,1fr);
                  gap: 16px; }}
      .v-card {{ background: #f9f9f9; padding: 12px;
                  border-radius: 8px; }}
      .v-card h3 {{ color: #006D77; margin-bottom: 8px;
                     font-size: 14px; }}
      .v-card video {{ width: 100%; border-radius: 8px; }}
      @media(max-width:600px) {{
        .ss-grid {{ grid-template-columns: repeat(2,1fr); }}
        .v-grid  {{ grid-template-columns: 1fr; }}
      }}
    </style>
  </head>
  <body>
    <div class="header">
      <h1>🏪 Super Business Shop</h1>
      <p>Mukammal Guide — {TIMESTAMP}</p>
      <p>🌐 https://super-business-flutter-web.onrender.com</p>
    </div>
  '''

    for sec in SECTIONS[1:]:
      html += f'''
    <div class="section">
      <h2>{sec["heading"]}</h2>
      <pre>{sec["content"].strip()}</pre>
    </div>'''

    html += f'''
    <div class="section">
      <h2>📸 Screenshots</h2>
      <div class="ss-grid">{ss_tags}</div>
    </div>
    <div class="section">
      <h2>🎥 Feature Videos</h2>
      <div class="v-grid">{video_tags}</div>
    </div>
  </body></html>'''

    with open('test/manual/user_manual.html', 'w',
              encoding='utf-8') as f:
      f.write(html)
    print('✅ HTML: test/manual/user_manual.html')

  make_docx()
  make_pdf()
  make_txt()
  make_html()
  print('\\n✅ All manuals saved to test/manual/')
